import io
import zipfile
from typing import Optional
from bs4 import BeautifulSoup
from flask import current_app
from app import firestore_db

class ExportService:
    @staticmethod
    def _resolve_project_id(project_id: Optional[str], chapter_id: str) -> str:
        if project_id:
            return project_id

        matches = []
        for chapter_doc in firestore_db.collection_group('chapters').stream():
            if chapter_doc.id != chapter_id:
                continue
            project_ref = chapter_doc.reference.parent.parent
            if project_ref:
                matches.append(project_ref.id)

        if not matches:
            raise ValueError("Chapter not found")
        if len(matches) > 1:
            raise ValueError("projectId is required because chapter_id is not unique")
        return matches[0]

    @staticmethod
    def _html_to_markdown(html_content: str) -> str:
        """Sederhana merubah HTML ke Markdown."""
        if not html_content:
            return ""
        soup = BeautifulSoup(html_content, "html.parser")
        return soup.get_text("\n\n", strip=True)

    @staticmethod
    def export_chapter_to_md(project_id: str, chapter_id: str) -> tuple[bytes, str]:
        """Return (file_bytes, filename)"""
        project_id = ExportService._resolve_project_id(project_id, chapter_id)
        chap_ref = firestore_db.collection('projects').document(project_id).collection('chapters').document(chapter_id).get()
        if not chap_ref.exists:
            raise ValueError("Chapter not found")
            
        data = chap_ref.to_dict()
        title = data.get('title', 'Untitled_Chapter')
        content = data.get('content', '')
        
        md_text = f"# {title}\n\n{ExportService._html_to_markdown(content)}"
        
        filename = f"{title.replace(' ', '_')}.md"
        return md_text.encode('utf-8'), filename

    @staticmethod
    def export_chapter_to_docx(project_id: str, chapter_id: str) -> tuple[bytes, str]:
        """Return (file_bytes, filename)"""
        project_id = ExportService._resolve_project_id(project_id, chapter_id)
        chap_ref = firestore_db.collection('projects').document(project_id).collection('chapters').document(chapter_id).get()
        if not chap_ref.exists:
            raise ValueError("Chapter not found")
            
        data = chap_ref.to_dict()
        title = data.get('title', 'Untitled_Chapter')
        content = data.get('content', '')
        
        filename = f"{title.replace(' ', '_')}.docx"
        
        try:
            from docx import Document
            doc = Document()
            doc.add_heading(title, 0)
            
            # Sangat sederhana: paragraf dari html to text
            md_text = ExportService._html_to_markdown(content)
            for paragraph in md_text.split("\n\n"):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
                    
            f = io.BytesIO()
            doc.save(f)
            f.seek(0)
            return f.read(), filename
        except ImportError:
            # Fallback jika python-docx tidak ada: return sebagai markdown text aja tapi extension .doc
            # Ini fallback sederhana.
            md_text = f"# {title}\n\n{ExportService._html_to_markdown(content)}"
            return md_text.encode('utf-8'), f"{title.replace(' ', '_')}.doc"

    @staticmethod
    def export_project_to_zip(project_id: str) -> tuple[bytes, str]:
        """Return (file_bytes, filename)"""
        proj_ref = firestore_db.collection('projects').document(project_id).get()
        if not proj_ref.exists:
            raise ValueError("Project not found")
            
        proj_data = proj_ref.to_dict()
        proj_title = proj_data.get('title', 'Project')
        
        chapters = firestore_db.collection('projects').document(project_id).collection('chapters').stream()
        
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
            for ch in chapters:
                data = ch.to_dict()
                title = data.get('title', 'Untitled')
                content = data.get('content', '')
                md_text = f"# {title}\n\n{ExportService._html_to_markdown(content)}"
                
                # Tambahkan ke zip
                filename = f"{data.get('index', 0):02d}_{title.replace(' ', '_')}.md"
                zf.writestr(filename, md_text.encode('utf-8'))
                
        zip_buffer.seek(0)
        zip_filename = f"{proj_title.replace(' ', '_')}_Export.zip"
        
        return zip_buffer.read(), zip_filename
