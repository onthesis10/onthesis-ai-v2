import requests
import urllib.parse
import logging
import json
import os
import uuid
import io
import datetime
from io import BytesIO
from dotenv import load_dotenv


# Load Environment Variables
load_dotenv()

from flask import Blueprint, render_template, request, jsonify, Response, stream_with_context, send_file, current_app, redirect, url_for
from flask_login import login_required, current_user
from firebase_admin import firestore
from pypdf import PdfReader
from bs4 import BeautifulSoup
from docx import Document

# Import module internal aplikasi
import app
from app import limiter
from app.services.ai_service import AIService
from app.utils import ai_utils
from app.utils.citation_helper import generate_bibliography

# Inisialisasi Blueprint dan Logger
from . import assistant_bp
# Inisialisasi Blueprint dan Logger
from . import assistant_bp
logger = logging.getLogger(__name__)

_agent_supervisor = None

LEGACY_WRITING_REMOVAL_PAYLOAD = {
    "error": "LEGACY_WRITING_ROUTE_REMOVED",
    "message": "Route writing lama sudah dinonaktifkan. Gunakan runtime blueprint baru melalui /api/agent/run atau halaman /writing.",
    "preferred_route": "/api/agent/run",
    "preferred_page": "/writing",
}


def _serialize_firestore_value(value):
    if hasattr(value, "isoformat"):
        try:
            return value.isoformat()
        except Exception:
            return str(value)
    if hasattr(value, "to_datetime"):
        try:
            return value.to_datetime().isoformat()
        except Exception:
            return str(value)
    return value


def _serialize_project_record(doc_id: str, data: dict) -> dict:
    created_at = data.get("createdAt") or data.get("created_at")
    updated_at = data.get("lastUpdated") or data.get("updatedAt") or data.get("updated_at")
    end_date = data.get("endDate") or data.get("end_date")
    return {
        "id": doc_id,
        "userId": data.get("userId"),
        "title": data.get("title", "Tanpa Judul"),
        "description": data.get("description", ""),
        "problem_statement": data.get("problem_statement", ""),
        "methodology": data.get("methodology", ""),
        "variables": data.get("variables", ""),
        "content": data.get("content", ""),
        "status": data.get("status", "DRAFT"),
        "progress": int(data.get("progress", 0) or 0),
        "createdAt": _serialize_firestore_value(created_at),
        "created_at": _serialize_firestore_value(created_at),
        "lastUpdated": _serialize_firestore_value(updated_at),
        "updatedAt": _serialize_firestore_value(updated_at),
        "updated_at": _serialize_firestore_value(updated_at),
        "endDate": _serialize_firestore_value(end_date),
        "end_date": _serialize_firestore_value(end_date),
    }


def _serialize_citation_record(doc_id: str, data: dict) -> dict:
    created_at = data.get("createdAt") or data.get("created_at")
    return {
        "id": doc_id,
        "projectId": data.get("projectId", ""),
        "userId": data.get("userId", ""),
        "type": data.get("type", ""),
        "title": data.get("title", ""),
        "author": data.get("author", ""),
        "year": str(data.get("year", "") or ""),
        "journal": data.get("journal", ""),
        "publisher": data.get("publisher", ""),
        "url": data.get("url", ""),
        "doi": data.get("doi", ""),
        "volume": data.get("volume", ""),
        "issue": data.get("issue", ""),
        "pages": data.get("pages", ""),
        "notes": data.get("notes", ""),
        "pdfUrl": data.get("pdfUrl", ""),
        "createdAt": _serialize_firestore_value(created_at),
        "created_at": _serialize_firestore_value(created_at),
    }


def _build_project_write_payload(payload: dict, user_id: str, include_created_at: bool = False) -> dict:
    now = firestore.SERVER_TIMESTAMP
    write_payload = {
        "userId": user_id,
        "title": str(payload.get("title") or "Proyek Baru").strip() or "Proyek Baru",
        "description": str(payload.get("description") or "").strip(),
        "problem_statement": str(payload.get("problem_statement") or "").strip(),
        "methodology": str(payload.get("methodology") or "").strip(),
        "variables": payload.get("variables") or "",
        "content": payload.get("content") or "",
        "status": str(payload.get("status") or "DRAFT").strip() or "DRAFT",
        "progress": int(payload.get("progress", 0) or 0),
        "updated_at": now,
        "updatedAt": now,
        "lastUpdated": now,
    }
    if payload.get("endDate") is not None:
        write_payload["endDate"] = payload.get("endDate")
    if include_created_at:
        write_payload["created_at"] = now
        write_payload["createdAt"] = now
    return write_payload

def get_agent_supervisor():
    global _agent_supervisor
    if _agent_supervisor is None:
        from app.agent.supervisor import SupervisorAgent
        _agent_supervisor = SupervisorAgent()
    return _agent_supervisor


def _legacy_writing_route_removed():
    return jsonify(dict(LEGACY_WRITING_REMOVAL_PAYLOAD)), 410


def _redirect_legacy_writing_page():
    project_id = request.args.get("id") or request.args.get("project_id")
    destination = url_for("main.writing")
    if project_id:
        destination = f"{destination}?id={urllib.parse.quote(str(project_id))}"
    return redirect(destination, code=302)


def _legacy_generator_task_to_agent_prompt(task_type: str, data: dict, thesis_context_str: str = "") -> str:
    """Legacy adapter kept for compatibility; new writing-agent flows should use `/api/agent/run`."""
    input_text = (data or {}).get("input_text", "")
    title = (data or {}).get("context_title", "")
    problem = (data or {}).get("context_problem", "")
    chapter = _detect_chapter_from_task(task_type)
    variable_name = (data or {}).get("variable_name", "")

    task_prompts = {
        "general": "Buatkan draft akademik yang relevan dengan konteks tesis dan editor saat ini.",
        "chat": input_text or "Jawab pertanyaan pengguna berdasarkan konteks tesis saat ini.",
        "continue": f"Lanjutkan paragraf akademik berikut dengan tetap konsisten terhadap konteks tesis:\n{input_text}",
        "improve": f"Perbaiki paragraf berikut agar lebih akademik, koheren, dan rapi:\n{input_text}",
        "paraphrase": f"Parafrase teks berikut tanpa mengubah makna akademiknya:\n{input_text}",
        "literature_review": "Susun literature review untuk Bab 2 berdasarkan konteks tesis dan referensi yang tersedia.",
        "bab1_latar_belakang": "Tulis latar belakang penelitian untuk Bab 1 berdasarkan konteks tesis saat ini.",
        "bab1_part_ideal": "Tulis paragraf kondisi ideal pada latar belakang Bab 1.",
        "bab1_part_factual": "Tulis paragraf kondisi faktual pada latar belakang Bab 1.",
        "bab1_part_gap": "Formulasikan research gap yang jelas untuk Bab 1.",
        "bab1_part_solution": "Tulis paragraf solusi atau arah penelitian untuk Bab 1.",
        "bab1_rumusan": "Susun rumusan masalah penelitian yang tajam untuk Bab 1.",
        "bab1_tujuan": "Susun tujuan penelitian yang selaras dengan rumusan masalah.",
        "bab2_kajian_pustaka": "Susun kajian pustaka Bab 2 yang terstruktur dan akademik.",
        "bab2_teori": "Jelaskan landasan teori utama untuk Bab 2 secara akademik.",
        "bab2_part_x": f"Jelaskan kajian teori untuk {variable_name or 'variabel X'} pada Bab 2.",
        "bab2_part_y": f"Jelaskan kajian teori untuk {variable_name or 'variabel Y'} pada Bab 2.",
        "bab2_part_context": "Jelaskan konteks penelitian yang relevan untuk Bab 2.",
        "bab2_part_relation": "Jelaskan hubungan antar variabel atau konsep pada Bab 2.",
        "bab2_part_framework": "Susun kerangka pemikiran penelitian untuk Bab 2.",
        "bab2_part_hypothesis": "Susun hipotesis penelitian yang logis untuk Bab 2.",
        "bab3_metode": "Tulis metodologi penelitian Bab 3 yang konsisten dengan konteks tesis.",
        "bab3_part_approach": "Jelaskan pendekatan penelitian pada Bab 3.",
        "bab3_part_loc": "Jelaskan lokasi penelitian pada Bab 3.",
        "bab3_part_pop": "Jelaskan populasi dan sampel penelitian pada Bab 3.",
        "bab3_part_var": "Jelaskan definisi variabel penelitian pada Bab 3.",
        "bab3_part_inst": "Jelaskan instrumen penelitian pada Bab 3.",
        "bab3_part_val": "Jelaskan uji validitas dan reliabilitas pada Bab 3.",
        "bab3_part_ana": "Jelaskan teknik analisis data pada Bab 3.",
        "bab3_part_proc": "Jelaskan prosedur penelitian pada Bab 3.",
        "methodology": "Susun penjelasan metodologi penelitian secara akademik untuk Bab 3.",
        "bab4_pembahasan": "Tulis hasil dan pembahasan Bab 4 berdasarkan konteks yang tersedia.",
        "bab4_part_descriptive": "Tulis analisis statistik deskriptif untuk Bab 4.",
        "bab4_part_discussion": "Tulis pembahasan hasil penelitian untuk Bab 4.",
        "bab4_part_implication": "Tulis implikasi hasil penelitian untuk Bab 4.",
        "bab4_part_object": "Jelaskan objek atau temuan utama penelitian pada Bab 4.",
        "bab4_part_qualitative": "Tulis interpretasi hasil kualitatif untuk Bab 4.",
        "bab4_part_prerequisite": "Tulis hasil uji prasyarat analisis untuk Bab 4.",
        "bab4_part_hypothesis": "Tulis hasil uji hipotesis untuk Bab 4.",
        "discussion_chapter4": "Susun pembahasan utama Bab 4 yang terhubung ke Bab 2.",
        "bab5_penutup": "Tulis bab penutup yang mencakup kesimpulan, implikasi, dan saran.",
        "bab5_part_conclusion": "Tulis kesimpulan penelitian untuk Bab 5.",
        "bab5_part_implication": "Tulis implikasi penelitian untuk Bab 5.",
        "bab5_part_suggestion": "Tulis saran penelitian untuk Bab 5.",
        "conclusion": "Susun kesimpulan penelitian secara ringkas dan akademik.",
        "validate_citations": "Periksa seluruh sitasi pada draft ini dan jelaskan masalah sitasi yang perlu diperbaiki.",
    }

    instruction = task_prompts.get(task_type)
    if not instruction:
        instruction = f"Bantu tulis bagian tesis untuk {chapter} berdasarkan konteks proyek saat ini."

    context_lines = []
    if title:
        context_lines.append(f"Judul: {title}")
    if problem:
        context_lines.append(f"Masalah: {problem}")
    if input_text and task_type not in {"chat", "continue", "improve", "paraphrase"}:
        context_lines.append(f"Catatan pengguna: {input_text}")
    if thesis_context_str:
        context_lines.append(f"Konteks Thesis Brain:\n{thesis_context_str}")

    if context_lines:
        return f"{instruction}\n\n" + "\n".join(context_lines)
    return instruction

# ==============================================================================
# BAGIAN 1: HELPER LIMITATION CHECKER (RATE LIMITING)
# ==============================================================================

def check_limits(user, limit_type='generator'):
    """
    Mengecek apakah user sudah mencapai batas penggunaan harian.
    
    Args:
        user: Objek current_user Flask-Login
        limit_type: 'generator' (Max 3x) atau 'chat' (Max 4x)
    
    Returns:
        (Allowed: bool, Message: str)
    """
    # 1. BYPASS UNTUK PRO USER
    # Jika user punya flag is_pro = True, maka bebas limit
    if getattr(user, 'is_pro', False):
        return True, "Pro User - Unlimited Access"

    # 2. CEK DATABASE UNTUK FREE USER
    # Gunakan tanggal hari ini sebagai key agar reset setiap hari
    today_str = datetime.datetime.now().strftime('%Y-%m-%d')
    
    # Gunakan nama koleksi yang berbeda agar kuota tidak tercampur
    # usage_logs_gen  -> Kuota Tools Berat (Generator, Paraphrase, Outline)
    # usage_logs_chat -> Kuota Chat Ringan
    collection_name = 'usage_logs_gen' if limit_type == 'generator' else 'usage_logs_chat'
    max_limit = 3 if limit_type == 'generator' else 4
    
    # ID Dokumen: UserID_Tanggal
    doc_ref = app.firestore_db.collection(collection_name).document(f"{user.id}_{today_str}")
    doc = doc_ref.get()

    current_usage = 0
    if doc.exists:
        current_usage = doc.to_dict().get('count', 0)

    # Logika Pengecekan
    if current_usage >= max_limit:
        return False, f"Kuota Harian {limit_type.capitalize()} Habis ({current_usage}/{max_limit}). Upgrade ke Pro untuk akses tanpa batas!"
    
    return True, "OK"

def increment_limit(user, limit_type='generator'):
    """
    Menambah hitungan penggunaan (+1) setelah request sukses.
    Hanya dijalankan untuk User Free.
    """
    if getattr(user, 'is_pro', False): 
        return

    today_str = datetime.datetime.now().strftime('%Y-%m-%d')
    collection_name = 'usage_logs_gen' if limit_type == 'generator' else 'usage_logs_chat'
    
    doc_ref = app.firestore_db.collection(collection_name).document(f"{user.id}_{today_str}")
    
    if doc_ref.get().exists:
        # Jika dokumen ada, increment count
        doc_ref.update({'count': firestore.Increment(1)})
    else:
        # Jika dokumen belum ada (request pertama hari ini), buat baru
        doc_ref.set({
            'count': 1, 
            'userId': user.id, 
            'date': today_str,
            'email': getattr(user, 'email', 'unknown')
        })


# ==============================================================================
# BAGIAN 2: HALAMAN VIEW (FRONTEND RENDERING)
# ==============================================================================

@assistant_bp.route('/writing-assistant')
@login_required
def writing_assistant():
    """Legacy page alias. Redirect permanently to the SPA Writing Studio."""
    return _redirect_legacy_writing_page()

@assistant_bp.route('/writing-studio')
@login_required
def writing_studio_page():
    """Legacy page alias. Redirect permanently to the SPA Writing Studio."""
    return _redirect_legacy_writing_page()

@assistant_bp.route('/generator-kajian-teori')
@login_required
def generator_kajian_teori():
    """Halaman khusus workbench Kajian Teori (Bab 2)."""
    return render_template('generator_kajian_teori.html')

@assistant_bp.route('/thesis-defense')
@login_required
def thesis_defense_page():
    """Halaman simulasi sidang skripsi."""
    return render_template('thesis_defense.html')


# ==============================================================================
# BAGIAN 3: API MANAJEMEN PROJECT (CRUD)
# ==============================================================================

@assistant_bp.route('/api/projects/new', methods=['POST'])
@login_required
def api_create_project():
    """Membuat project baru."""
    try:
        new_project = _build_project_write_payload({}, str(current_user.id), include_created_at=True)
        update_time, project_ref = app.firestore_db.collection('projects').add(new_project)
        return jsonify({'status': 'success', 'projectId': project_ref.id})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@assistant_bp.route('/api/projects', methods=['POST'])
@login_required
def create_project():
    """Create project melalui backend API canonical."""
    try:
        payload = request.get_json() or {}
        new_project = _build_project_write_payload(payload, str(current_user.id), include_created_at=True)
        _update_time, project_ref = app.firestore_db.collection('projects').add(new_project)
        snapshot = project_ref.get()
        data = snapshot.to_dict() if snapshot.exists else new_project
        return jsonify({
            'status': 'success',
            'projectId': project_ref.id,
            'project': _serialize_project_record(project_ref.id, data),
        }), 201
    except Exception as e:
        logger.error(f"Create Project Error: {e}")
        return jsonify({'status': 'error', 'message': str(e)}), 500

@assistant_bp.route('/api/projects', methods=['GET'])
@login_required
def get_user_projects():
    """Mengambil daftar project milik user."""
    try:
        docs = app.firestore_db.collection('projects')\
            .where('userId', '==', str(current_user.id))\
            .stream()
            
        projects = []
        for doc in docs:
            d = doc.to_dict()
            projects.append(_serialize_project_record(doc.id, d))
        
        # Sort client-side karena firestore composite index kadang ribet
        projects.sort(key=lambda x: str(x.get('lastUpdated') or x.get('updated_at') or ''), reverse=True)
            
        return jsonify({'status': 'success', 'projects': projects})
    except Exception as e:
        logger.error(f"List Projects Error: {e}")
        return jsonify({'status': 'error', 'projects': [], 'message': str(e)}), 500

@assistant_bp.route('/api/projects/<project_id>', methods=['GET'])
@login_required
def get_project_details(project_id):
    """Mengambil detail satu project."""
    try:
        doc_ref = app.firestore_db.collection('projects').document(project_id)
        doc = doc_ref.get()
        
        if not doc.exists:
            return jsonify({'error': 'Project not found'}), 404
            
        data = doc.to_dict()
        if data.get('userId') != str(current_user.id):
            return jsonify({'error': 'Unauthorized'}), 403
            
        return jsonify({'status': 'success', 'project': _serialize_project_record(doc.id, data)})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@assistant_bp.route('/api/projects/<project_id>', methods=['PUT'])
@login_required
def update_project(project_id):
    """Update data project (Auto-save)."""
    try:
        data = request.get_json()
        doc_ref = app.firestore_db.collection('projects').document(project_id)
        doc = doc_ref.get()
        
        if not doc.exists:
            return jsonify({'error': 'Project not found'}), 404
        if doc.to_dict().get('userId') != str(current_user.id):
            return jsonify({'error': 'Unauthorized'}), 403
            
        write_payload = _build_project_write_payload(data, str(current_user.id), include_created_at=False)
        doc_ref.set(write_payload, merge=True)
        
        updated_doc = doc_ref.get()
        return jsonify({
            'status': 'success',
            'project': _serialize_project_record(project_id, updated_doc.to_dict() if updated_doc.exists else write_payload),
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@assistant_bp.route('/api/projects/<project_id>', methods=['DELETE'])
@login_required
def delete_project(project_id):
    """Delete project dan referensi terkait via backend API."""
    try:
        doc_ref = app.firestore_db.collection('projects').document(project_id)
        doc = doc_ref.get()
        if not doc.exists:
            return jsonify({'error': 'Project not found'}), 404
        if doc.to_dict().get('userId') != str(current_user.id):
            return jsonify({'error': 'Unauthorized'}), 403

        citations = (
            app.firestore_db.collection('citations')
            .where('projectId', '==', project_id)
            .where('userId', '==', str(current_user.id))
            .stream()
        )
        for citation_doc in citations:
            citation_doc.reference.delete()

        doc_ref.delete()
        return jsonify({'status': 'success'}), 200
    except Exception as e:
        logger.error(f"Delete Project Error: {e}")
        return jsonify({'status': 'error', 'message': str(e)}), 500


# ==============================================================================
# BAGIAN 4: API REFERENCES & CITATIONS
# ==============================================================================

@assistant_bp.route('/api/citations', methods=['GET'])
@login_required
def get_user_citations():
    """List semua citation milik user, opsional filter per project."""
    try:
        project_id = request.args.get('projectId', '').strip()
        query_ref = app.firestore_db.collection('citations').where('userId', '==', str(current_user.id))
        if project_id:
            query_ref = query_ref.where('projectId', '==', project_id)

        docs = query_ref.stream()
        citations = [_serialize_citation_record(doc.id, doc.to_dict() or {}) for doc in docs]
        citations.sort(key=lambda x: str(x.get('createdAt') or ''), reverse=True)
        return jsonify({'status': 'success', 'citations': citations}), 200
    except Exception as e:
        logger.error(f"List Citations Error: {e}")
        return jsonify({'status': 'error', 'citations': [], 'message': str(e)}), 500


@assistant_bp.route('/api/citations', methods=['POST'])
@login_required
def create_citation():
    """Create citation lewat backend canonical API."""
    try:
        ref_data = request.get_json() or {}
        project_id = str(ref_data.get('projectId') or '').strip()
        if not project_id:
            return jsonify({'status': 'error', 'message': 'projectId wajib diisi'}), 400

        proj_ref = app.firestore_db.collection('projects').document(project_id)
        proj = proj_ref.get()
        if not proj.exists or proj.to_dict().get('userId') != str(current_user.id):
            return jsonify({'status': 'error', 'message': 'Unauthorized'}), 403

        ref_data['projectId'] = project_id
        ref_data['userId'] = str(current_user.id)
        ref_data['createdAt'] = firestore.SERVER_TIMESTAMP
        ref_data.pop('id', None)

        _update_time, doc_ref = app.firestore_db.collection('citations').add(ref_data)
        doc_ref.update({'id': doc_ref.id})
        snapshot = doc_ref.get()
        return jsonify({
            'status': 'success',
            'message': 'Referensi tersimpan',
            'id': doc_ref.id,
            'citation': _serialize_citation_record(doc_ref.id, snapshot.to_dict() if snapshot.exists else ref_data),
        }), 201
    except Exception as e:
        logger.error(f"Create Citation Error: {e}")
        return jsonify({'error': str(e)}), 500

@assistant_bp.route('/api/project/<project_id>/references/add', methods=['POST'])
@login_required
def add_project_reference(project_id):
    """Menambah referensi ke project (Safe Mode dengan .add)."""
    try:
        ref_data = request.get_json()
        if not ref_data:
            return jsonify({'status': 'error', 'message': 'Data kosong'}), 400

        # Verifikasi kepemilikan project
        proj_ref = app.firestore_db.collection('projects').document(project_id)
        proj = proj_ref.get()
        if not proj.exists or proj.to_dict().get('userId') != str(current_user.id):
            return jsonify({'status': 'error', 'message': 'Unauthorized'}), 403

        # Inject Data Penting
        ref_data['projectId'] = project_id
        ref_data['userId'] = str(current_user.id)
        ref_data['createdAt'] = firestore.SERVER_TIMESTAMP
        
        # Bersihkan field ID jika ada (biar Firestore yang generate)
        if 'id' in ref_data:
            del ref_data['id']

        # Simpan ke Firestore
        update_time, doc_ref = app.firestore_db.collection('citations').add(ref_data)
        
        # Update dokumen agar punya field 'id' yang sama dengan doc ID (opsional tapi berguna untuk frontend)
        doc_ref.update({'id': doc_ref.id})

        return jsonify({'status': 'success', 'message': 'Referensi tersimpan', 'id': doc_ref.id}), 200

    except Exception as e:
        logger.error(f"Add Reference Error: {e}")
        return jsonify({'error': str(e)}), 500


@assistant_bp.route('/api/citations/<citation_id>', methods=['DELETE'])
@login_required
def delete_citation(citation_id):
    """Delete citation lewat backend canonical API."""
    try:
        doc_ref = app.firestore_db.collection('citations').document(citation_id)
        doc = doc_ref.get()
        if not doc.exists:
            return jsonify({'error': 'Not found'}), 404
        if doc.to_dict().get('userId') != str(current_user.id):
            return jsonify({'error': 'Unauthorized'}), 403
        doc_ref.delete()
        return jsonify({'status': 'success'}), 200
    except Exception as e:
        logger.error(f"Delete Citation Error: {e}")
        return jsonify({'error': str(e)}), 500

@assistant_bp.route('/api/references/delete', methods=['POST'])
@login_required
def delete_reference():
    """Menghapus referensi."""
    try:
        data = request.get_json()
        ref_id = data.get('id')
        
        if not ref_id:
            return jsonify({'error': 'No Reference ID'}), 400
            
        doc_ref = app.firestore_db.collection('citations').document(ref_id)
        doc = doc_ref.get()
        
        if doc.exists and doc.to_dict().get('userId') == str(current_user.id):
            doc_ref.delete()
            return jsonify({'status': 'success'})
        else:
            return jsonify({'error': 'Not found or Unauthorized'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ==============================================================================
# BAGIAN 5: API AI GENERATOR UTAMA (STREAMING)
# ==============================================================================

@assistant_bp.route('/api/writing-assistant', methods=['POST'])
@login_required
@limiter.limit("13 per minute")
def api_writing_assistant():
    return _legacy_writing_route_removed()


# ==============================================================================
# BAGIAN 6: API CHAT (STREAMING)
# ==============================================================================

@assistant_bp.route('/chat/stream', methods=['POST'])
@login_required
@limiter.limit("13 per minute")
def chat_with_ai_stream():
    return _legacy_writing_route_removed()

# Endpoint Chat Alias (untuk kompatibilitas frontend lama)
# Endpoint Chat Alias (Hybrid: Page vs API)
@assistant_bp.route('/chat', methods=['GET', 'POST'])
@login_required
def chat_hybrid_route(): 
    if request.method == 'GET':
        return _redirect_legacy_writing_page()
    return _legacy_writing_route_removed()

@assistant_bp.route('/api/assistant/chat/copilot', methods=['POST'])
@login_required
@limiter.limit("13 per minute")
def chat_copilot():
    """
    Endpoint for Context-Aware Research Co-Pilot with Memory.
    """
    try:
        logger.info(f"Copilot Request from: {current_user.id} | Valid: {current_user.is_authenticated}")
        
        # 1. Limit Check
        allowed, msg = check_limits(current_user, 'chat')
        if not allowed:
            return jsonify({"error": "LIMIT_REACHED", "message": msg}), 403

        data = request.get_json()
        if not data:
             logger.error("Copilot Error: No JSON Body")
             return jsonify({'error': 'No JSON body'}), 400
             
        message = data.get('message')
        research_context = data.get('researchContext', data.get('research_context', {}))
        data_context = data.get('data_context', {})
        analysis_state = data.get('analysis_state', {})
        conversation_history = data.get('conversation_history', [])

        logger.info(f"Processing Copilot Message: {message[:50]}... | History: {len(conversation_history)} msgs")

        # 2. Call Service with History
        response = AIService.chat_research_copilot(
            current_user, 
            message, 
            research_context, 
            data_context, 
            analysis_state,
            conversation_history
        )
        
        # 3. Usage Tracking
        increment_limit(current_user, 'chat')

        return jsonify(response)
    except Exception as e:
        logger.error(f"Chat Copilot Error: {e}")
        return jsonify({'error': str(e)}), 500


# ==============================================================================
# BAGIAN 7: API TOOLS SPESIFIK (GENERATOR TOOLS)
# ==============================================================================

@assistant_bp.route('/api/generate-outline', methods=['POST'])
@login_required
@limiter.limit("13 per minute")
def api_generate_outline():
    return _legacy_writing_route_removed()

@assistant_bp.route('/api/paraphrase', methods=['POST'])
@login_required
@limiter.limit("13 per minute")
def paraphrase():
    return _legacy_writing_route_removed()

@assistant_bp.route('/expand-text', methods=['POST'])
@login_required
@limiter.limit("13 per minute")
def expand_text_endpoint():
    return _legacy_writing_route_removed()

@assistant_bp.route('/api/ai/edit-text', methods=['POST'])
@login_required
@limiter.limit("13 per minute")
def ai_edit_text():
    return _legacy_writing_route_removed()


# ==============================================================================
# BAGIAN 8: API ANALYSIS & LOGIC (HEAVY TOOLS)
# ==============================================================================

@assistant_bp.route('/api/analyze-style', methods=['POST'])
@login_required
def api_analyze_style():
    """Menganalisis gaya penulisan dari file user."""
    if 'file' not in request.files: return jsonify({'error': 'No file'}), 400
    try:
        # Style analysis biasanya task sekali jalan, hitungannya generator
        allowed, msg = check_limits(current_user, 'generator')
        if not allowed: return jsonify({'error': 'LIMIT_REACHED', 'message': msg}), 403
        
        file = request.files['file']
        # Membaca file dan analisa (Logic ada di AIService)
        profile = AIService.analyze_style_from_file(current_user, file)
        
        increment_limit(current_user, 'generator')
        return jsonify({'status': 'success', 'style_profile': profile})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@assistant_bp.route('/api/logic-matrix', methods=['POST'])
@login_required
def logic_matrix_endpoint():
    """Membuat Matrix Konsistensi Logika (Masalah vs Kesimpulan)."""
    try:
        allowed, msg = check_limits(current_user, 'generator')
        if not allowed: return jsonify({'error': 'LIMIT_REACHED', 'message': msg}), 403

        data = request.get_json()
        problem = data.get('problem', '').strip()
        conclusion = data.get('conclusion', '').strip()
        
        if not problem or not conclusion:
            return jsonify({'status': 'error', 'message': 'Data kurang lengkap.'}), 400
            
        result = AIService.generate_logic_matrix(current_user, problem, conclusion)
        
        increment_limit(current_user, 'generator')
        return jsonify({'status': 'success', 'data': result})
        
    except Exception as e:
        logger.error(f"Logic Matrix Error: {e}")
        return jsonify({'status': 'error', 'message': str(e)}), 500

@assistant_bp.route('/api/logic/check', methods=['POST'])
@login_required
def check_logic():
    """Logic Check Menyeluruh (Bab 1 vs Bab 5)."""
    try:
        allowed, msg = check_limits(current_user, 'generator')
        if not allowed: return jsonify({'error': 'LIMIT_REACHED', 'message': msg}), 403

        data = request.get_json()
        # AIService.check_logic_consistency biasanya return JSON
        result = AIService.check_logic_consistency(data)
        
        increment_limit(current_user, 'generator')
        return jsonify({'status': 'success', 'data': result})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500

@assistant_bp.route('/api/check-method-compliance', methods=['POST'])
@login_required
def check_method_compliance_endpoint():
    """Cek apakah teks sesuai dengan metodologi (Kuali/Kuanti)."""
    try:
        # Ini fitur analisis, kita bisa masukkan ke generator limit atau free
        # Mari masukkan ke generator limit
        allowed, msg = check_limits(current_user, 'generator')
        if not allowed: return jsonify({'error': 'LIMIT_REACHED', 'message': msg}), 403

        data = request.get_json()
        text = data.get('text', '').strip()
        method_mode = data.get('method_mode', 'quantitative') 
        
        if not text: return jsonify({'status': 'error'}), 400

        issues = AIService.check_method_compliance(text, method_mode)
        
        increment_limit(current_user, 'generator')
        return jsonify({'status': 'success', 'issues': issues, 'checked_method': method_mode})
        
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500


# ==============================================================================
# BAGIAN 9: API DEFENSE & PPT (ADVANCED)
# ==============================================================================

@assistant_bp.route('/api/defense/<action>', methods=['POST'])
@login_required
def defense_endpoint(action):
    """
    Endpoint Simulasi Sidang.
    Action: 'start', 'answer', 'evaluate'.
    """
    try:
        # Simulasi sidang sangat berat, wajib cek limit
        allowed, msg = check_limits(current_user, 'generator')
        if not allowed and action == 'start': # Cek limit hanya saat mulai sesi
             return jsonify({'error': 'LIMIT_REACHED', 'message': msg}), 403

        data = request.get_json()
        result = AIService.thesis_defense_simulation(current_user, action, data)
        
        # Increment limit jika memulai sesi baru
        if action == 'start':
            increment_limit(current_user, 'generator')

        response_key = 'response' if action != 'evaluate' else 'report'
        return jsonify({'status': 'success', response_key: result})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@assistant_bp.route('/api/generate-ppt', methods=['POST'])
@login_required
def generate_ppt_endpoint():
    """Generate Powerpoint Slide."""
    try:
        allowed, msg = check_limits(current_user, 'generator')
        if not allowed: return jsonify({'error': 'LIMIT_REACHED', 'message': msg}), 403

        pptx_file = AIService.generate_ppt(current_user, request.get_json())
        filename = f"OnThesis_Slide_{datetime.datetime.now().strftime('%Y%m%d')}.pptx"
        
        increment_limit(current_user, 'generator')
        
        return send_file(
            pptx_file,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.presentationml.presentation'
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ==============================================================================
# BAGIAN 10: API UTILS (SEARCH, BIBLIOGRAPHY, EXPORT)
# ==============================================================================

@assistant_bp.route('/api/extract-pdf-simple', methods=['POST'])
@login_required
def extract_pdf_simple():
    """Ekstrak teks dari PDF upload."""
    try:
        if 'document' not in request.files:
            return jsonify({'status': 'error', 'error': 'No file'}), 400
            
        file = request.files['document']
        reader = PdfReader(file)
        full_text = []
        max_pages = min(len(reader.pages), 50) 
        
        for i in range(max_pages):
            page_text = reader.pages[i].extract_text()
            if page_text: full_text.append(page_text)
        
        joined_text = "\n".join(full_text)
        
        if len(joined_text.strip()) < 50:
             return jsonify({'status': 'error', 'error': 'Teks tidak terbaca (Scan?).'}), 422

        return jsonify({
            'status': 'success', 
            'data': {'filename': file.filename, 'content': joined_text, 'page_count': len(reader.pages)}
        })
    except Exception as e:
        return jsonify({'status': 'error', 'error': str(e)}), 500

@assistant_bp.route('/api/unified-search-references', methods=['POST'])
@login_required
def unified_search_references():
    """
    Search references dari Crossref/DOAJ/dll.
    Biasanya fitur ini gratis agar user mudah cari referensi.
    """
    try:
        data = request.get_json()
        query = data.get('query', '').strip()
        year_filter = str(data.get('year', '')).strip()
        
        if not (year_filter.isdigit() and len(year_filter) == 4): year_filter = None
        if not query: return jsonify({'status': 'error', 'message': 'Query kosong'}), 400

        HEADERS = {'User-Agent': 'Mozilla/5.0 (OnThesis Academic Bot)'}
        results = []

        # Logic search sederhana ke Crossref (Implementasi penuh bisa via Helper)
        params = {'query': query, 'rows': 5}
        if year_filter: params['filter'] = f'from-pub-date:{year_filter}'
        
        try:
            resp = requests.get("https://api.crossref.org/works", params=params, headers=HEADERS, timeout=5)
            if resp.status_code == 200:
                items = resp.json().get('message', {}).get('items', [])
                for item in items:
                    try:
                        title = item.get('title', [''])[0]
                        auths = item.get('author', [])
                        authors_str = ", ".join([f"{a.get('family', '')}" for a in auths[:2]])
                        pub_date = item.get('published-print', {}).get('date-parts', [[None]])[0][0]
                        
                        results.append({
                            'title': title, 
                            'author': authors_str, 
                            'source': 'Crossref', 
                            'year': str(pub_date) if pub_date else 'n.d.',
                            'doi': item.get('DOI', ''),
                            'url': item.get('URL', '')
                        })
                    except: continue
        except Exception as api_err:
            logger.warning(f"Search API Error: {api_err}")

        return jsonify({'status': 'success', 'results': results}) 

    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500

@assistant_bp.route('/api/generate-bibliography', methods=['POST'])
@login_required
def api_generate_bibliography():
    """Generate Daftar Pustaka otomatis."""
    try:
        data = request.get_json()
        project_id = data.get('projectId')
        if not project_id: return jsonify({'error': 'No Project ID'}), 400
        
        refs_query = app.firestore_db.collection('citations').where('projectId', '==', project_id).stream()
        references = [doc.to_dict() for doc in refs_query]
        
        # Panggil helper function
        html_out, text_out = generate_bibliography(references)
        
        return jsonify({'status': 'success', 'html': html_out, 'text': text_out})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@assistant_bp.route('/api/export-docx', methods=['POST'])
@login_required
def export_docx_endpoint():
    """Export Project ke DOCX."""
    try:
        data = request.get_json()
        project_id = data.get('projectId')
        
        doc_ref = app.firestore_db.collection('projects').document(project_id)
        doc = doc_ref.get()
        if not doc.exists: return jsonify({'error': 'Not found'}), 404
        
        project_data = doc.to_dict()
        html_content = project_data.get('content', '')
        title = project_data.get('title', 'Draft Skripsi')

        # Setup Document
        document = Document()
        document.add_heading(title, 0)
        
        author_name = getattr(current_user, 'name', None) or getattr(current_user, 'username', 'Mahasiswa')
        document.add_paragraph(f"Penulis: {author_name}")
        document.add_paragraph(f"Generated by OnThesis AI")
        document.add_page_break()

        # Parse HTML content (Simple Version)
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Iterasi elemen penting saja
        for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'li']):
            text = element.get_text().strip()
            if not text: continue
            
            if element.name == 'h1': document.add_heading(text, level=1)
            elif element.name == 'h2': document.add_heading(text, level=2)
            elif element.name == 'h3': document.add_heading(text, level=3)
            elif element.name == 'li': document.add_paragraph(text, style='List Paragraph')
            else: document.add_paragraph(text)

        file_stream = BytesIO()
        document.save(file_stream)
        file_stream.seek(0)
        
        safe_title = "".join([c for c in title if c.isalnum() or c in (' ', '_')]).rstrip()
        filename = f"{safe_title[:30]}_Draft.docx"

        return send_file(
            file_stream, 
            as_attachment=True, 
            download_name=filename, 
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@assistant_bp.route('/api/review-document', methods=['POST'])
def review_document():
    """
    Endpoint Review Document (AI Editor).
    Note: Biasanya diakses oleh plugin editor.
    """
    try:
        # Cek limit jika mau, tapi request ini dari plugin editor
        # Kita anggap generator usage
        if current_user.is_authenticated:
            allowed, msg = check_limits(current_user, 'generator')
            if not allowed: return jsonify({'status': 'error', 'message': msg}), 403
            increment_limit(current_user, 'generator')

        data = request.get_json() or {}
        text = data.get("text", "").strip()
        
        if not text:
            return jsonify({"error": "Tidak ada teks"}), 400

        # Panggil logika review (bisa via AIService atau langsung litellm)
        # Disini kita mock implementasi sederhana atau panggil service
        # Agar konsisten, kita panggil logic review dari AIService jika ada,
        # atau implementasi inline jika service belum siap.
        
        # Implementasi inline untuk review grammar/typo
        # (Idealnya dipindah ke AIService)
        from app.utils.ai_utils import safe_completion as completion
        
        prompt = """
        You are a professional academic reviewer.
        TASK: Review the text for grammar errors, typo, and clarity issues.
        OUTPUT: JSON Array [{"target": "bad text", "issue": "explanation", "fix": "correction"}].
        NO MARKDOWN.
        """
        
        resp = completion(
            model="groq/llama-3.3-70b-versatile",
            messages=[{"role":"system", "content": prompt}, {"role":"user", "content": text}],
            response_format={"type": "json_object"}
        )
        
        content = resp.choices[0].message.content
        result_json = json.loads(ai_utils.clean_json_output(content))
        
        # Handle format beda dari LLM
        reviews = result_json if isinstance(result_json, list) else result_json.get('reviews', [])
        
        return jsonify({"status": "success", "reviews": reviews})

    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@assistant_bp.route('/api/assistant/logic-check', methods=['POST'])
@login_required
def logic_check_route():
    """
    ENDPOINT: AUDIT KONSISTENSI (BENANG MERAH).
    Mengecek keselarasan antara Judul vs Masalah vs Tujuan.
    """
    try:
        # 1. CEK LIMIT (Fitur ini "Berat", masuk kuota generator)
        allowed, msg = check_limits(current_user, 'generator')
        if not allowed:
            return jsonify({'error': 'LIMIT_REACHED', 'message': msg}), 403

        data = request.get_json()
        
        # Validasi Input Minimal
        if not data.get('title') or not data.get('problem'):
            return jsonify({'error': 'Judul dan Masalah wajib diisi untuk audit.'}), 400

        # 2. PANGGIL AI SERVICE
        result = AIService.analyze_logic_flow(current_user, data)
        
        # 3. INCREMENT LIMIT (Jika sukses)
        increment_limit(current_user, 'generator')
        
        return jsonify(result)

    except Exception as e:
        logger.error(f"Logic Check Error: {e}")
        return jsonify({'error': str(e)}), 500

# ==============================================================================
# BAGIAN 10: API AI GENERATOR UTAMA (STREAMING)
# ==============================================================================
@assistant_bp.route('/api/assistant/generate-stream', methods=['POST'])
@login_required
def generate_stream_endpoint():
    return _legacy_writing_route_removed()


@assistant_bp.route('/api/orchestrator/execute', methods=['POST'])
@login_required
def orchestrator_execute():
    return _legacy_writing_route_removed()

def _detect_chapter_from_task(task_type: str) -> str:
    """Map task_type to chapter for context compilation"""
    task_chapter_map = {
        # Bab 1
        'bab1_part_ideal': 'bab1',
        'bab1_part_factual': 'bab1',
        'bab1_part_gap': 'bab1',
        'bab1_part_solution': 'bab1',
        'bab1_rumusan': 'bab1',
        'bab1_tujuan': 'bab1',
        'bab1_latar_belakang': 'bab1',
        # Bab 2
        'bab2_kajian_pustaka': 'bab2',
        'bab2_teori': 'bab2',
        'bab2_part_x': 'bab2',
        'bab2_part_y': 'bab2',
        'bab2_part_context': 'bab2',
        'bab2_part_relation': 'bab2',
        'bab2_part_framework': 'bab2',
        'bab2_part_hypothesis': 'bab2',
        'literature_review': 'bab2',
        # Bab 3
        'bab3_metode': 'bab3',
        'bab3_part_approach': 'bab3',
        'bab3_part_loc': 'bab3',
        'bab3_part_pop': 'bab3',
        'bab3_part_var': 'bab3',
        'bab3_part_inst': 'bab3',
        'bab3_part_val': 'bab3',
        'bab3_part_ana': 'bab3',
        'bab3_part_proc': 'bab3',
        'methodology': 'bab3',
        # Bab 4
        'bab4_pembahasan': 'bab4',
        'bab4_part_descriptive': 'bab4',
        'bab4_part_discussion': 'bab4',
        'bab4_part_implication': 'bab4',
        'bab4_part_object': 'bab4',
        'bab4_part_qualitative': 'bab4',
        'bab4_part_prerequisite': 'bab4',
        'bab4_part_hypothesis': 'bab4',
        'discussion_chapter4': 'bab4',
        # Bab 5
        'bab5_penutup': 'bab5',
        'bab5_part_conclusion': 'bab5',
        'bab5_part_implication': 'bab5',
        'bab5_part_suggestion': 'bab5',
        'conclusion': 'bab5',
    }
    return task_chapter_map.get(task_type, 'bab1')
