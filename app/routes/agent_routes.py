from flask import Blueprint, request, jsonify, Response, stream_with_context
from app.services.agent_service import agent_app
from langchain_core.messages import HumanMessage
import json
import os
import uuid
import time

agent_bp = Blueprint('agent_bp', __name__)

TEMP_DIR = os.path.join(os.getcwd(), 'temp_data')
os.makedirs(TEMP_DIR, exist_ok=True)

@agent_bp.route('/api/agent/upload', methods=['POST'])
def upload_file():
    """
    Upload a CSV file for analysis.
    Returns the file path to be used in the chat session.
    """
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    # Generate a unique session ID if not provided, or use one
    session_id = request.form.get('session_id', str(uuid.uuid4()))
    
    # Create session directory
    session_dir = os.path.join(TEMP_DIR, session_id)
    os.makedirs(session_dir, exist_ok=True)
    
    # Save file
    filename = file.filename
    file_path = os.path.join(session_dir, filename)
    file.save(file_path)
    
    return jsonify({
        'status': 'success',
        'session_id': session_id,
        'file_path': file_path,
        'filename': filename
    })

@agent_bp.route('/api/agent/stream', methods=['POST'])
def stream_chat_agent():
    """
    Stream the agent's execution process.
    """
    data = request.json
    user_message = data.get('message')
    dataset_path = data.get('dataset_path')
    
    if not user_message:
        return jsonify({'error': 'Message is required'}), 400

    def generate():
        # Context to send to the graph
        inputs = {
            "messages": [HumanMessage(content=user_message)],
            "dataset_path": dataset_path
        }
        
        # Stream events from the graph
        # "updates" mode streams the state updates from each node
        for event in agent_app.stream(inputs):
            for node_name, node_state in event.items():
                
                # 1. Notify STEP START/UPDATE
                if node_name == 'planner':
                    route = node_state.get('route')
                    content = node_state['messages'][-1].content
                    
                    if route == 'conversation':
                        # If conversation, this IS the response. Stream as response type.
                        yield f"data: {json.dumps({'type': 'response', 'content': content})}\n\n"
                    else:
                        # If analysis, stream as a step
                        yield f"data: {json.dumps({'type': 'step', 'id': 'plan', 'title': 'Planning Analysis', 'status': 'completed', 'content': content})}\n\n"
                
                elif node_name == 'coder':
                    code = node_state.get('code_generated', '')
                    yield f"data: {json.dumps({'type': 'step', 'id': 'code', 'title': 'Generating Python Code', 'status': 'completed', 'content': code})}\n\n"
                
                elif node_name == 'executor':
                    result = node_state.get('execution_result', '')
                    error = node_state.get('error')
                    status = 'failed' if error else 'completed'
                    artifacts = node_state.get('artifacts', [])
                    
                    yield f"data: {json.dumps({'type': 'step', 'id': 'exec', 'title': 'Executing Code', 'status': status, 'output': result})}\n\n"
                    
                    # Stream Artifacts
                    if artifacts:
                        for artifact in artifacts:
                             yield f"data: {json.dumps({'type': 'artifact', 'content': artifact})}\n\n"

                elif node_name == 'responder':
                    # Stream the final synthesized response text
                    resp_messages = node_state.get('messages', [])
                    if resp_messages:
                        final_text = resp_messages[-1].content
                        yield f"data: {json.dumps({'type': 'response', 'content': final_text})}\n\n"

        yield f"data: {json.dumps({'type': 'done'})}\n\n"

    return Response(stream_with_context(generate()), mimetype='text/event-stream')


@agent_bp.route('/api/agent/export/pdf', methods=['POST'])
def export_agent_pdf():
    """
    Export agent analysis results to PDF.
    Expects JSON: { title, content (markdown), artifacts: [{type, data}], dataset? }
    """
    import sys
    import markdown as md
    from datetime import datetime

    try:
        data = request.json
        title = data.get('title', 'Laporan Analisis')
        content_md = data.get('content', '')
        artifacts = data.get('artifacts', [])
        dataset_name = data.get('dataset', '')

        # Convert markdown to HTML
        content_html = md.markdown(
            content_md,
            extensions=['tables', 'fenced_code', 'nl2br']
        )

        # Extract image artifacts
        charts = [a for a in artifacts if a.get('type') == 'image_base64']

        # Build template data
        template_data = {
            'title': title,
            'content_html': content_html,
            'charts': charts,
            'date': datetime.now().strftime('%d %B %Y'),
            'dataset': dataset_name
        }

        # Use existing PdfService
        from app.services.pdf_service import PdfService
        pdf_bytes = PdfService.generate_pdf({'template': 'agent_report', **template_data})

        from flask import Response as FlaskResponse
        return FlaskResponse(
            pdf_bytes,
            status=200,
            mimetype='application/pdf',
            headers={
                'Content-Disposition': f'attachment; filename=OnThesis_Analysis_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf',
                'Content-Length': str(len(pdf_bytes)),
                'Cache-Control': 'no-cache',
                'Access-Control-Allow-Origin': '*'
            }
        )
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'status': 'error', 'message': str(e)}), 500


@agent_bp.route('/api/agent/export/docx', methods=['POST'])
def export_agent_docx():
    """
    Export agent analysis results to DOCX.
    Expects JSON: { title, content (markdown), artifacts: [{type, data}], dataset? }
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO
    import base64
    import re
    from datetime import datetime
    from flask import send_file

    try:
        data = request.json
        title = data.get('title', 'Laporan Analisis')
        content_md = data.get('content', '')
        artifacts = data.get('artifacts', [])
        dataset_name = data.get('dataset', '')

        doc = Document()

        # --- Styles ---
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        # --- Title ---
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0x0c, 0x4a, 0x6e)

        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f'Generated by OnThesis AI — {datetime.now().strftime("%d %B %Y")}')
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
        if dataset_name:
            run2 = subtitle.add_run(f' | Dataset: {dataset_name}')
            run2.font.size = Pt(10)
            run2.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

        doc.add_paragraph('')  # Spacer

        # --- Parse Markdown Content ---
        lines = content_md.split('\n')
        current_table_lines = []
        in_table = False
        in_code_block = False
        code_block_content = []

        for line in lines:
            stripped = line.strip()

            # Handle code blocks
            if stripped.startswith('```'):
                if in_code_block:
                    # End code block
                    code_text = '\n'.join(code_block_content)
                    p = doc.add_paragraph()
                    run = p.add_run(code_text)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    code_block_content = []
                    in_code_block = False
                else:
                    in_code_block = True
                continue

            if in_code_block:
                code_block_content.append(line)
                continue

            # Handle tables
            if '|' in stripped and stripped.startswith('|'):
                if re.match(r'^\|[\s\-:|]+\|$', stripped):
                    continue  # Skip separator row
                current_table_lines.append(stripped)
                in_table = True
                continue
            elif in_table:
                # End of table, flush it
                _add_table_to_docx(doc, current_table_lines)
                current_table_lines = []
                in_table = False

            # Headers
            if stripped.startswith('### '):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith('## '):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith('# '):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith('- ') or stripped.startswith('* '):
                text = stripped[2:]
                text = _clean_markdown_inline(text)
                doc.add_paragraph(text, style='List Bullet')
            elif re.match(r'^\d+\.\s', stripped):
                text = re.sub(r'^\d+\.\s', '', stripped)
                text = _clean_markdown_inline(text)
                doc.add_paragraph(text, style='List Number')
            elif stripped == '---' or stripped == '***':
                # Horizontal rule - add a thin paragraph
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
            elif stripped:
                text = _clean_markdown_inline(stripped)
                p = doc.add_paragraph(text)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Flush remaining table
        if current_table_lines:
            _add_table_to_docx(doc, current_table_lines)

        # --- Add Chart Images ---
        charts = [a for a in artifacts if a.get('type') == 'image_base64']
        if charts:
            doc.add_heading('Visualisasi Data', level=2)
            for i, chart in enumerate(charts):
                try:
                    img_data = base64.b64decode(chart['data'])
                    img_stream = BytesIO(img_data)
                    doc.add_picture(img_stream, width=Inches(5.5))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    caption = doc.add_paragraph(f'Gambar {i + 1}. Visualisasi Hasil Analisis')
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in caption.runs:
                        run.font.size = Pt(9)
                        run.font.italic = True
                        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
                except Exception as img_err:
                    print(f"DOCX: Error adding image {i}: {img_err}")
                    doc.add_paragraph(f'[Chart {i + 1} - gagal dimuat]')

        # --- Footer ---
        doc.add_paragraph('')
        footer_p = doc.add_paragraph()
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_p.add_run('Dokumen ini dihasilkan secara otomatis oleh OnThesis AI Analysis Assistant.')
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

        # Save to BytesIO
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        filename = f'OnThesis_Analysis_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        return send_file(
            file_stream,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'status': 'error', 'message': str(e)}), 500


def _clean_markdown_inline(text):
    """Remove markdown inline formatting (bold, italic, code) for plain DOCX text."""
    import re
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # Bold
    text = re.sub(r'\*(.+?)\*', r'\1', text)  # Italic
    text = re.sub(r'`(.+?)`', r'\1', text)  # Inline code
    return text


def _add_table_to_docx(doc, table_lines):
    """Parse markdown table lines and add to DOCX document."""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import re

    if not table_lines:
        return

    # Parse cells
    rows_data = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip('|').split('|')]
        cells = [_clean_markdown_inline(c) for c in cells]
        rows_data.append(cells)

    if not rows_data:
        return

    num_cols = len(rows_data[0])
    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    table.style = 'Table Grid'

    for i, row_data in enumerate(rows_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = row.cells[j]
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        run.font.name = 'Arial'

        # Style header row
        if i == 0:
            for j in range(num_cols):
                cell = row.cells[j]
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), '0284c7')
                cell._tc.get_or_add_tcPr().append(shading)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
